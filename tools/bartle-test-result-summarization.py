"""
Bartle Test of Gamer Psychology — Class Summary Report Generator

Usage:
    python bartle_report.py [path/to/bartle_data.csv]

Reads a CSV of person Bartle test results and generates a Word document
(bartle_report.docx) with summary statistics, graphs, and a results table.

CSV format:
    person,primary_type,primary_pct,type2,type2_pct,type3,type3_pct,type4,type4_pct
"""

import sys
import os
import tempfile
from datetime import date

import pandas as pd
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT


BARTLE_TYPES = ["Achiever", "Explorer", "Killer", "Socialiser"]
TYPE_COLORS = {
    "Achiever": "#4CAF50",
    "Explorer": "#2196F3",
    "Killer": "#F44336",
    "Socialiser": "#FF9800",
}


def load_data(csv_path):
    df = pd.read_csv(csv_path)
    expected = [
        "person", "primary_type", "primary_pct",
        "type2", "type2_pct", "type3", "type3_pct", "type4", "type4_pct",
    ]
    missing = [c for c in expected if c not in df.columns]
    if missing:
        raise ValueError(f"CSV missing columns: {', '.join(missing)}")
    return df


def get_person_type_scores(df):
    """Build a DataFrame with one row per person and columns for each Bartle type score."""
    rows = []
    for _, r in df.iterrows():
        scores = {r["primary_type"]: r["primary_pct"]}
        for i in [2, 3, 4]:
            t = r[f"type{i}"]
            p = r[f"type{i}_pct"]
            if pd.notna(t) and pd.notna(p):
                scores[str(t)] = p
        rows.append({"person": r["person"], **scores})
    result = pd.DataFrame(rows).fillna(0)
    for t in BARTLE_TYPES:
        if t not in result.columns:
            result[t] = 0
    return result


def compute_stats(df, scores_df):
    stats = {}
    total = len(df)
    stats["total"] = total

    # Primary type counts
    type_counts = df["primary_type"].value_counts()
    stats["type_counts"] = {t: int(type_counts.get(t, 0)) for t in BARTLE_TYPES}
    stats["type_pcts"] = {t: round(100 * stats["type_counts"][t] / total, 1) for t in BARTLE_TYPES}

    # Most/least common
    most_common = type_counts.idxmax()
    least_common_count = type_counts.min()
    least_common = [t for t in BARTLE_TYPES if type_counts.get(t, 0) == least_common_count]
    stats["most_common"] = most_common
    stats["most_common_count"] = int(type_counts.max())
    stats["least_common"] = least_common
    stats["least_common_count"] = int(least_common_count)

    # Class averages across all persons for each type
    stats["class_avg"] = {t: round(scores_df[t].mean(), 1) for t in BARTLE_TYPES}

    # Highest individual score
    max_score = df["primary_pct"].max()
    top_persons = df[df["primary_pct"] == max_score]
    stats["highest_score"] = int(max_score)
    stats["highest_persons"] = list(top_persons["person"])
    stats["highest_type"] = list(top_persons["primary_type"])

    # Lowest primary score
    min_score = df["primary_pct"].min()
    stats["lowest_score"] = int(min_score)

    return stats


def make_bar_chart(stats, path):
    types = BARTLE_TYPES
    counts = [stats["type_counts"][t] for t in types]
    colors = [TYPE_COLORS[t] for t in types]

    fig, ax = plt.subplots(figsize=(7, 4))
    bars = ax.bar(types, counts, color=colors, edgecolor="white", linewidth=1.2)
    for bar, count in zip(bars, counts):
        ax.text(bar.get_x() + bar.get_width() / 2, bar.get_height() + 0.15,
                str(count), ha="center", va="bottom", fontweight="bold", fontsize=12)
    ax.set_ylabel("Number of persons", fontsize=11)
    ax.set_title("Primary Bartle Type Distribution", fontsize=13, fontweight="bold")
    ax.set_ylim(0, max(counts) + 2)
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)
    plt.tight_layout()
    fig.savefig(path, dpi=150, bbox_inches="tight")
    plt.close(fig)


def make_pie_chart(stats, path):
    types = [t for t in BARTLE_TYPES if stats["type_counts"][t] > 0]
    counts = [stats["type_counts"][t] for t in types]
    colors = [TYPE_COLORS[t] for t in types]

    fig, ax = plt.subplots(figsize=(6, 5))
    wedges, texts, autotexts = ax.pie(
        counts, labels=types, colors=colors, autopct="%1.0f%%",
        startangle=90, pctdistance=0.75, textprops={"fontsize": 11},
    )
    for at in autotexts:
        at.set_fontweight("bold")
        at.set_color("white")
    ax.set_title("Primary Type Breakdown", fontsize=13, fontweight="bold")
    plt.tight_layout()
    fig.savefig(path, dpi=150, bbox_inches="tight")
    plt.close(fig)


def make_grouped_bar_chart(stats, path):
    types = BARTLE_TYPES
    avgs = [stats["class_avg"][t] for t in types]
    colors = [TYPE_COLORS[t] for t in types]

    fig, ax = plt.subplots(figsize=(7, 4))
    bars = ax.bar(types, avgs, color=colors, edgecolor="white", linewidth=1.2)
    for bar, avg in zip(bars, avgs):
        ax.text(bar.get_x() + bar.get_width() / 2, bar.get_height() + 0.5,
                f"{avg}%", ha="center", va="bottom", fontweight="bold", fontsize=11)
    ax.set_ylabel("Average Score (%)", fontsize=11)
    ax.set_title("Class Average Scores by Type", fontsize=13, fontweight="bold")
    ax.set_ylim(0, max(avgs) + 10)
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)
    plt.tight_layout()
    fig.savefig(path, dpi=150, bbox_inches="tight")
    plt.close(fig)


def build_docx(df, scores_df, stats, chart_paths, output_path):
    doc = Document()

    # -- Styles --
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # -- Title --
    title = doc.add_heading("Bartle Test of Gamer Psychology", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Class Summary Report")
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(100, 100, 100)

    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = date_para.add_run(f"Generated: {date.today().strftime('%B %d, %Y')}")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(140, 140, 140)

    doc.add_paragraph()

    # -- Summary Section --
    doc.add_heading("Summary", level=1)

    total = stats["total"]
    bullets = [
        f"Total persons surveyed: {total}",
    ]

    # Type breakdown
    for t in BARTLE_TYPES:
        c = stats["type_counts"][t]
        p = stats["type_pcts"][t]
        bullets.append(f"{t}: {c} person{'s' if c != 1 else ''} ({p}%)")

    bullets.append(
        f"Most common primary type: {stats['most_common']} "
        f"({stats['most_common_count']} persons)"
    )

    if stats["least_common_count"] == 0:
        least_str = ", ".join(stats["least_common"])
        bullets.append(f"No persons had a primary type of: {least_str}")
    else:
        least_str = ", ".join(stats["least_common"])
        bullets.append(
            f"Least common primary type: {least_str} "
            f"({stats['least_common_count']} person{'s' if stats['least_common_count'] != 1 else ''})"
        )

    bullets.append("")
    bullets.append("Class average scores across all types:")
    for t in BARTLE_TYPES:
        bullets.append(f"  {t}: {stats['class_avg'][t]}%")

    bullets.append("")
    top_names = ", ".join(stats["highest_persons"])
    top_types = ", ".join(stats["highest_type"])
    bullets.append(
        f"Highest individual primary score: {stats['highest_score']}% "
        f"({top_names} — {top_types})"
    )
    bullets.append(f"Lowest individual primary score: {stats['lowest_score']}%")

    for b in bullets:
        if b == "":
            doc.add_paragraph()
        elif b.startswith("  "):
            p = doc.add_paragraph(b.strip(), style="List Bullet 2")
        else:
            p = doc.add_paragraph(b, style="List Bullet")

    doc.add_paragraph()

    # -- Charts --
    doc.add_heading("Charts", level=1)

    for chart_path, caption in chart_paths:
        doc.add_picture(chart_path, width=Inches(5.5))
        last_para = doc.paragraphs[-1]
        last_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cap.add_run(caption)
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(120, 120, 120)
        run.font.italic = True
        doc.add_paragraph()

    # -- person Results Table --
    doc.add_heading("person Results", level=1)

    # Sort by primary type then by primary_pct descending
    sorted_df = df.sort_values(
        ["primary_type", "primary_pct"], ascending=[True, False]
    ).reset_index(drop=True)

    table = doc.add_table(rows=1, cols=6)
    table.style = "Medium Shading 1 Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    headers = ["person", "Primary Type", "Primary %", "2nd Type & %", "3rd Type & %", "4th Type & %"]
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(9)

    for _, r in sorted_df.iterrows():
        row = table.add_row()
        row.cells[0].text = str(r["person"])
        row.cells[1].text = str(r["primary_type"])
        row.cells[2].text = f"{int(r['primary_pct'])}%"
        row.cells[3].text = f"{r['type2']} ({int(r['type2_pct'])}%)"
        row.cells[4].text = f"{r['type3']} ({int(r['type3_pct'])}%)"
        row.cells[5].text = f"{r['type4']} ({int(r['type4_pct'])}%)"
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)

    doc.save(output_path)


def main():
    if len(sys.argv) > 1:
        csv_path = sys.argv[1]
    else:
        script_dir = os.path.dirname(os.path.abspath(__file__))
        csv_path = os.path.join(script_dir, "bartle_data.csv")

    if not os.path.exists(csv_path):
        print(f"Error: CSV file not found: {csv_path}")
        sys.exit(1)

    output_dir = os.path.dirname(os.path.abspath(csv_path))
    output_path = os.path.join(output_dir, "bartle_report.docx")

    print(f"Reading data from: {csv_path}")
    df = load_data(csv_path)
    print(f"Found {len(df)} persons.")

    scores_df = get_person_type_scores(df)
    stats = compute_stats(df, scores_df)

    # Generate charts to temp files
    tmp_dir = tempfile.mkdtemp()
    bar_path = os.path.join(tmp_dir, "bar_chart.png")
    pie_path = os.path.join(tmp_dir, "pie_chart.png")
    grouped_path = os.path.join(tmp_dir, "grouped_bar.png")

    print("Generating charts...")
    make_bar_chart(stats, bar_path)
    make_pie_chart(stats, pie_path)
    make_grouped_bar_chart(stats, grouped_path)

    chart_paths = [
        (bar_path, "Figure 1: Number of persons per primary Bartle type"),
        (pie_path, "Figure 2: Percentage breakdown of primary types"),
        (grouped_path, "Figure 3: Class average scores for each Bartle type"),
    ]

    print("Building report...")
    build_docx(df, scores_df, stats, chart_paths, output_path)

    # Cleanup temp files
    for p, _ in chart_paths:
        try:
            os.remove(p)
        except OSError:
            pass
    try:
        os.rmdir(tmp_dir)
    except OSError:
        pass

    print(f"Report saved to: {output_path}")


if __name__ == "__main__":
    main()
