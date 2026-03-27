"""
Securly Reports Weekly Summary Generator

Reads Securly Classroom session summary emails from Outlook (classic/2016),
aggregates browsing data for the current business week, and generates
per-class Word document reports.

Requirements:
    pip install pywin32 beautifulsoup4 python-docx lxml

Usage:
    1. Open Outlook 2016 (classic) — NOT the new Outlook app
    2. Run: python securly_report.py
    3. Check Desktop for generated .docx files
"""

import datetime
import os
import re
import sys
from collections import defaultdict

import win32com.client
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

# Domains exempt from referral qualification
REFERRAL_EXEMPT_DOMAINS = ["vex.com", "instructure.com", "office.com", "code.org"]

# ---------------------------------------------------------------------------
# Constants
# ---------------------------------------------------------------------------

CLASS_NAME_MAP = {
    "STEM-ADV ROBOTICS & AUTOMATION": "Adv. Robotics",
    "CS DISCOVERIES 2- GAME DESIGN": "Game Design",
    "DESIGN & MODELING": "Design & Modeling",
}

# Matches section codes like " - KDM078/01" at end of class name
SECTION_CODE_RE = re.compile(r"\s*-\s*[A-Z]{1,6}\d{1,4}/\d{1,3}\s*$")

DESKTOP_PATH = os.path.join(
    os.environ.get("USERPROFILE", ""),
    "OneDrive - issaquah.wednet.edu",
    "Desktop",
)

TOTAL_INSTRUCTIONAL_MINUTES = 235


# ---------------------------------------------------------------------------
# Date helpers
# ---------------------------------------------------------------------------

def get_week_range():
    """Return (monday, end_date) for the current business week."""
    today = datetime.date.today()
    monday = today - datetime.timedelta(days=today.weekday())
    # Cap at Friday if run on weekend
    friday = monday + datetime.timedelta(days=4)
    end = min(today, friday)
    return monday, end


# ---------------------------------------------------------------------------
# Outlook email retrieval
# ---------------------------------------------------------------------------

def fetch_emails(monday, end_date):
    """
    Connect to Outlook via COM and return a list of (HTMLBody, ReceivedTime)
    tuples from Inbox > Securly Reports within the date range.
    """
    try:
        outlook = win32com.client.Dispatch("Outlook.Application")
        ns = outlook.GetNamespace("MAPI")
    except Exception as e:
        print(f"ERROR: Could not connect to Outlook. Make sure Outlook 2016 (classic) is open.\n{e}")
        sys.exit(1)

    try:
        inbox = ns.GetDefaultFolder(6)  # 6 = olFolderInbox
    except Exception as e:
        print(f"ERROR: Could not access Inbox. Is Outlook connected to your account?\n{e}")
        sys.exit(1)

    try:
        folder = inbox.Folders["Securly Reports"]
    except Exception:
        print("ERROR: Could not find 'Securly Reports' subfolder in Inbox.")
        print("Available subfolders:")
        for i in range(inbox.Folders.Count):
            print(f"  - {inbox.Folders.Item(i + 1).Name}")
        sys.exit(1)

    # Build date filter (inclusive of end_date by using < end_date+1)
    start_str = monday.strftime("%m/%d/%Y")
    end_str = (end_date + datetime.timedelta(days=1)).strftime("%m/%d/%Y")
    filter_str = f"[ReceivedTime] >= '{start_str}' AND [ReceivedTime] < '{end_str}'"

    items = folder.Items.Restrict(filter_str)

    results = []
    for item in items:
        try:
            results.append((item.HTMLBody, item.ReceivedTime))
        except Exception:
            continue  # skip non-mail items

    return results


# ---------------------------------------------------------------------------
# HTML parsing
# ---------------------------------------------------------------------------

def clean_class_name(raw_name):
    """Strip section code and map to display name."""
    # Remove section code like " - KDM078/01"
    stripped = SECTION_CODE_RE.sub("", raw_name).strip()
    # Look up in map, fall back to title case
    return CLASS_NAME_MAP.get(stripped, stripped.title())


def parse_email_html(html):
    """
    Parse a Securly session summary email and return a list of records:
    [{"class_name": str, "date": str, "student": str, "website": str, "minutes": float}, ...]
    """
    soup = BeautifulSoup(html, "lxml")
    records = []

    # --- Extract class name and date from SUMMARY section ---
    class_name = None
    email_date = None

    # Find the cell containing class name (large font, ~30px, in SUMMARY area)
    for td in soup.find_all("td"):
        style = td.get("style", "")
        text = td.get_text(strip=True)
        if "font-size: 30px" in style or "font-size:30px" in style:
            if text:
                class_name = clean_class_name(text)
                break

    # Find the date (MM/DD/YYYY pattern)
    date_pattern = re.compile(r"\d{2}/\d{2}/\d{4}")
    for td in soup.find_all("td"):
        text = td.get_text(strip=True)
        m = date_pattern.match(text)
        if m:
            email_date = m.group(0)
            break

    if not class_name or not email_date:
        return records  # skip unparseable emails

    # --- Extract student browsing history ---
    # Each student block is in a bordered div containing a table
    # Student name is in a bold td (font-weight: 600, font-size: 16px)
    # within the STUDENT BROWSING HISTORY section

    # Find all bordered divs (student blocks)
    student_divs = soup.find_all("div", style=lambda s: s and "border-radius: 10px" in s and "border: 1px" in s)

    for div in student_divs:
        # Find student name - bold td with specific styling
        student_name = None
        name_tds = div.find_all("td", style=lambda s: s and "font-weight: 600" in s and "font-size: 16px" in s)
        for name_td in name_tds:
            text = name_td.get_text(strip=True)
            if text and "," in text:  # "Last, First" format
                student_name = text
                break

        if not student_name:
            continue

        # Find website rows within this student block
        # Each website row has: favicon img, anchor with domain, time, minutes bar+span
        # The minutes are in spans with font-weight: 600 followed by "min"
        tables_in_div = div.find_all("table", style=lambda s: s and "padding-left: 24px" in s)

        for tbl in tables_in_div:
            # Extract website domain
            website = None
            anchor = tbl.find("a")
            if anchor:
                website = anchor.get_text(strip=True)
            else:
                # Try favicon img domain
                img = tbl.find("img")
                if img:
                    src = img.get("src", "")
                    domain_match = re.search(r"domain=([^&\"]+)", src)
                    if domain_match:
                        website = domain_match.group(1)

            if not website:
                continue

            # Extract minutes
            minutes = 0.0
            bold_spans = tbl.find_all("span", style=lambda s: s and "font-weight: 600" in s)
            for span in bold_spans:
                # Check if next sibling contains "min"
                next_sib = span.find_next_sibling("span")
                if next_sib and "min" in next_sib.get_text(strip=True).lower():
                    try:
                        minutes = float(span.get_text(strip=True))
                    except ValueError:
                        pass
                    break

            records.append({
                "class_name": class_name,
                "date": email_date,
                "student": student_name,
                "website": website,
                "minutes": minutes,
            })

    return records


# ---------------------------------------------------------------------------
# Aggregation
# ---------------------------------------------------------------------------

def aggregate(records):
    """
    Aggregate records into:
    {class_name: {(student, website): {"minutes": float, "dates": set}}}
    """
    data = defaultdict(lambda: defaultdict(lambda: {"minutes": 0.0, "dates": set()}))
    for rec in records:
        key = (rec["student"], rec["website"])
        entry = data[rec["class_name"]][key]
        entry["minutes"] += rec["minutes"]
        entry["dates"].add(rec["date"])
    return data


# ---------------------------------------------------------------------------
# Deduction / referral helpers
# ---------------------------------------------------------------------------

def calc_deduction(minutes):
    if minutes < 10:
        return "0"
    elif minutes < 20:
        return "-1"
    else:
        return "-2"


def is_referral_exempt(website):
    """Check if a website matches any referral-exempt domain."""
    w = website.lower()
    return any(domain in w for domain in REFERRAL_EXEMPT_DOMAINS)


def calc_referral(minutes, website):
    if is_referral_exempt(website):
        return "-"
    return "Yes" if minutes >= 10 else "-"


def calc_instructional_pct(minutes):
    pct = (minutes / TOTAL_INSTRUCTIONAL_MINUTES) * 100
    return f"{pct:.1f}%"


# ---------------------------------------------------------------------------
# Word document generation
# ---------------------------------------------------------------------------

def set_cell_shading(cell, color):
    """Apply background shading to a table cell."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def set_font_arial(run, size=Pt(8)):
    """Set a run's font to Arial at the given size."""
    run.font.name = "Arial"
    run.font.size = size


def generate_docx(class_name, student_data, monday, end_date):
    """Generate a Word document for one class."""
    doc = Document()

    # Set default font to Arial for the whole document
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(8)

    # Set narrow margins
    for section in doc.sections:
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    # Title
    title = doc.add_paragraph()
    run = title.add_run(f"{class_name} - Week of {monday.strftime('%B %d')} - {end_date.strftime('%B %d, %Y')}")
    run.bold = True
    set_font_arial(run, Pt(14))
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Organize by student
    students = defaultdict(list)
    for (student, website), info in student_data.items():
        students[student].append((website, info))

    # Sort students alphabetically, websites alphabetically within each
    sorted_students = sorted(students.items())
    for student, sites in sorted_students:
        sites.sort(key=lambda x: x[0].lower())

    # Count total rows
    total_rows = sum(len(sites) for _, sites in sorted_students)

    # Create table
    cols = [
        "Student Name",
        "Websites Visited",
        "Dates Visited",
        "Total Minutes\nVisited",
        "Participation\nDeduction",
        "Referral\nQualified",
        "Total Instructional\nMinutes Spent",
        "Total Instructional\n%",
    ]
    table = doc.add_table(rows=1 + total_rows, cols=len(cols))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    hdr = table.rows[0]
    for i, col_name in enumerate(cols):
        cell = hdr.cells[i]
        cell.text = col_name
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True
                set_font_arial(run, Pt(8))
        set_cell_shading(cell, "D9E2F3")

    # Data rows
    row_idx = 1
    for student, sites in sorted_students:
        first_row = row_idx
        for site_idx, (website, info) in enumerate(sites):
            row = table.rows[row_idx]
            minutes = info["minutes"]
            dates_sorted = sorted(info["dates"])

            # Format dates as MM/DD only
            dates_display = []
            for d in dates_sorted:
                # Input dates are MM/DD/YYYY
                parts = d.split("/")
                if len(parts) == 3:
                    dates_display.append(f"{parts[0]}/{parts[1]}")
                else:
                    dates_display.append(d)

            # Format minutes: drop trailing zero (e.g. 12.0 -> 12, 12.5 -> 12.5)
            mins_str = f"{minutes:g}"

            # Only show student name on the first row for that student
            student_display = student if site_idx == 0 else ""

            values = [
                student_display,
                website,
                ", ".join(dates_display),
                f"{minutes:.1f}",
                calc_deduction(minutes),
                calc_referral(minutes, website),
                mins_str,
                calc_instructional_pct(minutes),
            ]

            for col, val in enumerate(values):
                cell = row.cells[col]
                cell.text = val
                for paragraph in cell.paragraphs:
                    for r in paragraph.runs:
                        set_font_arial(r, Pt(8))

            row_idx += 1

        # Merge student name cells for the same student
        if row_idx - first_row > 1:
            merged = table.cell(first_row, 0).merge(table.cell(row_idx - 1, 0))
            # Ensure merged cell keeps the name with correct font
            merged.text = student
            for paragraph in merged.paragraphs:
                for r in paragraph.runs:
                    set_font_arial(r, Pt(8))

    # Set column widths (approximate)
    col_widths = [1.1, 1.3, 0.9, 0.7, 0.85, 0.65, 0.85, 0.7]
    for row in table.rows:
        for i, width in enumerate(col_widths):
            row.cells[i].width = Inches(width)

    # Save
    filename = f"{class_name}-Week of {monday.isoformat()}-{end_date.isoformat()}.docx"
    filepath = os.path.join(DESKTOP_PATH, filename)
    doc.save(filepath)
    print(f"  Saved: {filename}")
    return filepath


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def main():
    monday, end_date = get_week_range()
    print(f"Processing week: {monday} to {end_date}")
    print()

    # Fetch emails from Outlook
    email_data = fetch_emails(monday, end_date)
    if not email_data:
        print("No emails found in 'Securly Reports' for this date range.")
        print("Make sure Outlook 2016 (classic) is open and connected.")
        return

    print(f"Found {len(email_data)} emails")

    # Parse all emails
    all_records = []
    for html, received_time in email_data:
        records = parse_email_html(html)
        all_records.extend(records)

    if not all_records:
        print("No student browsing data found in any emails.")
        return

    print(f"Parsed {len(all_records)} browsing records")
    print()

    # Aggregate
    aggregated = aggregate(all_records)
    print(f"Found {len(aggregated)} class(es):")
    for cn in sorted(aggregated.keys()):
        students = set(s for s, w in aggregated[cn].keys())
        print(f"  {cn}: {len(students)} students, {len(aggregated[cn])} website entries")
    print()

    # Generate Word docs
    print("Generating reports...")
    for class_name, student_data in sorted(aggregated.items()):
        generate_docx(class_name, student_data, monday, end_date)

    print()
    print("Done! Check your Desktop for the generated .docx files.")


if __name__ == "__main__":
    main()
